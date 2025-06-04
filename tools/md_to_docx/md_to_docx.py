import io
import logging
import re
from typing import Generator

import markdown
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from htmldocx import HtmlToDocx

from tools.utils.file_utils import get_meta_data
from tools.utils.mimetype_utils import MimeType
from tools.utils.param_utils import get_md_text


class MarkdownToDocxTool(Tool):
    logger = logging.getLogger(__name__)

    def _invoke(self, tool_parameters: dict) -> Generator[ToolInvokeMessage, None, None]:
        """
        invoke tools
        """
        # get parameters
        md_text = get_md_text(tool_parameters, is_strip_wrapper=True)
        try:
            # Legacy: using markdowntodocx lib
            # with NamedTemporaryFile(suffix=".docx", delete=True) as temp_docx_file:
            #     markdownconverter.markdownToWordFromString(string=md_text, outfile=temp_docx_file)
            #     result_file_bytes = Path(temp_docx_file.name).read_bytes()

            html = markdown.markdown(text=md_text, extensions=["extra", "toc"])

            # transform html to docx
            new_parser = HtmlToDocx()
            doc: Document = new_parser.parse_html_string(html)

            # Set fonts for all text elements
            self.set_fonts_for_all_runs(doc)

            result_bytes_io = io.BytesIO()
            doc.save(result_bytes_io)
            result_file_bytes = result_bytes_io.getvalue()
        except Exception as e:
            self.logger.exception("Failed to convert file")
            yield self.create_text_message(f"Failed to convert markdown text to DOCX file, error: {str(e)}")
            return

        yield self.create_blob_message(
            blob=result_file_bytes,
            meta=get_meta_data(
                mime_type=MimeType.DOCX,
                output_filename=tool_parameters.get("output_filename"),
            ),
        )
        return

    def is_contains_chinese_chars(self, text: str) -> bool:
        return bool(re.search(r'[\u4e00-\u9fff]', text))

    def set_chinese_fonts(self, doc):
        # Setting fonts globally
        # https://github.com/python-openxml/python-docx/issues/346#issuecomment-1698885586
        # https://zhuanlan.zhihu.com/p/548039429
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        rPr = style.element.get_or_add_rPr()
        rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def set_fonts_for_all_runs(self, doc: Document):
        """Set Times New Roman for English text and SimSun for Chinese text in all text elements."""

        # Process all paragraphs in the document
        paragraph: Paragraph
        for paragraph in doc.paragraphs:
            run: Run
            for run in paragraph.runs:
                self.apply_fonts_to_run(run)

        # Process all paragraphs in tables
        table: Table
        for table in doc.tables:
            for row in table.rows:
                cell: _Cell
                for cell in row.cells:
                    paragraph: Paragraph
                    for paragraph in cell.paragraphs:
                        run: Run
                        for run in paragraph.runs:
                            self.apply_fonts_to_run(run)

    def apply_fonts_to_run(self, run: Run):
        if not run or not run.text or not run.text.strip():  # Skip empty text
            return
        # Set default font to Times New Roman
        run.font.name = 'Times New Roman'
        # Set East Asian font to SimSun
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # Set ASCII font to Times New Roman
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        # Set high ANSI font to Times New Roman
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
